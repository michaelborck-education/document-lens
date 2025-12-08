"""
Document processing service for text extraction from various file formats
"""

import io
import json
import re
from typing import Any

from fastapi import HTTPException

# Import libraries for different file types
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None  # type: ignore[assignment]
    PdfReader = None  # type: ignore[misc,assignment]

try:
    from pdfplumber.pdf import PDF as PDFPlumberPDF  # noqa: N811
except ImportError:
    PDFPlumberPDF = None  # type: ignore[assignment, misc]

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[assignment]

try:
    from pptx import Presentation
except ImportError:
    Presentation = None


class DocumentProcessor:
    """Handles text extraction from various document formats"""

    def __init__(self) -> None:
        self.pdf_available = PdfReader is not None
        self.docx_available = Document is not None
        self.pptx_available = Presentation is not None

    async def extract_text(self, content: bytes, content_type: str, filename: str) -> str:
        """
        Extract text from document based on content type

        Args:
            content: Raw file content as bytes
            content_type: MIME type of the file
            filename: Name of the file (for error reporting)

        Returns:
            Extracted text as string
        """
        try:
            if content_type == "application/pdf":
                return await self._extract_from_pdf(content, filename)
            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return await self._extract_from_docx(content, filename)
            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.presentationml.presentation"
            ):
                raise ValueError(
                    "PPTX files are no longer supported in DocumentLens. "
                    "Please use PresentationLens service for presentation analysis. "
                    "See: https://github.com/michael-borck/presentation-lens"
                )
            elif content_type in ["text/plain", "text/markdown"]:
                return await self._extract_from_text(content, filename)
            elif content_type == "application/json":
                return await self._extract_from_json(content, filename)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

        except Exception as e:
            raise HTTPException(
                status_code=422, detail=f"Failed to extract text from {filename}: {e!s}"
            ) from e

    async def extract_text_with_pages(
        self, content: bytes, content_type: str, filename: str
    ) -> dict[str, Any]:
        """
        Extract text from document with page-level granularity.

        Args:
            content: Raw file content as bytes
            content_type: MIME type of the file
            filename: Name of the file (for error reporting)

        Returns:
            Dictionary with full_text, pages list, and total_pages
        """
        try:
            if content_type == "application/pdf":
                return await self._extract_from_pdf_with_pages(content, filename)
            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return await self._extract_from_docx_with_pages(content, filename)
            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.presentationml.presentation"
            ):
                raise ValueError(
                    "PPTX files are no longer supported in DocumentLens. "
                    "Please use PresentationLens service for presentation analysis. "
                    "See: https://github.com/michael-borck/presentation-lens"
                )
            elif content_type in ["text/plain", "text/markdown"]:
                text = await self._extract_from_text(content, filename)
                # Text files are treated as single page
                return {
                    "full_text": text,
                    "pages": [{"page_number": 1, "text": text}],
                    "total_pages": 1,
                }
            elif content_type == "application/json":
                text = await self._extract_from_json(content, filename)
                return {
                    "full_text": text,
                    "pages": [{"page_number": 1, "text": text}],
                    "total_pages": 1,
                }
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

        except Exception as e:
            raise HTTPException(
                status_code=422, detail=f"Failed to extract text from {filename}: {e!s}"
            ) from e

    async def _extract_from_pdf(self, content: bytes, filename: str) -> str:
        """Extract text from PDF file"""
        result = await self._extract_from_pdf_with_pages(content, filename)
        full_text: str = result["full_text"]
        return full_text

    async def _extract_from_pdf_with_pages(self, content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from PDF file with page-level granularity"""
        if not PyPDF2:
            raise ImportError("PyPDF2 not available. Please install with: pip install PyPDF2")

        pages: list[dict[str, Any]] = []

        try:
            # Try with PyPDF2 first
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    pages.append({"page_number": page_num, "text": text})
                else:
                    # Keep empty pages to maintain page numbering
                    pages.append({"page_number": page_num, "text": ""})

            # If PyPDF2 didn't extract much text, try pdfplumber
            total_text = "".join(p["text"] for p in pages)
            if len(total_text) < 100 and PDFPlumberPDF is not None:
                pdf_file.seek(0)
                with PDFPlumberPDF(pdf_file) as pdf:
                    pages = []
                    for page_num, page in enumerate(pdf.pages, start=1):  # type: ignore[assignment]
                        text = page.extract_text()
                        if text:
                            pages.append({"page_number": page_num, "text": text})
                        else:
                            pages.append({"page_number": page_num, "text": ""})

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e!s}") from e

        # Filter out completely empty pages for full_text, but keep page structure
        non_empty_texts = [p["text"] for p in pages if p["text"].strip()]
        if not non_empty_texts:
            raise ValueError("No text could be extracted from PDF")

        return {
            "full_text": "\n\n".join(non_empty_texts),
            "pages": pages,
            "total_pages": len(pages),
        }

    async def _extract_from_docx(self, content: bytes, filename: str) -> str:
        """Extract text from DOCX file"""
        result = await self._extract_from_docx_with_pages(content, filename)
        full_text: str = result["full_text"]
        return full_text

    async def _extract_from_docx_with_pages(self, content: bytes, filename: str) -> dict[str, Any]:
        """
        Extract text from DOCX file with page-level granularity.

        Note: DOCX files don't have fixed pages like PDFs. We treat the entire
        document as a single page unless page breaks are detected.
        """
        if Document is None:
            raise ImportError(
                "python-docx not available. Please install with: pip install python-docx"
            )

        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e!s}") from e

        if not text_parts:
            raise ValueError("No text could be extracted from DOCX")

        full_text = "\n\n".join(text_parts)

        # DOCX doesn't have native page structure like PDF
        # Return as single page (desktop app can handle differently if needed)
        return {
            "full_text": full_text,
            "pages": [{"page_number": 1, "text": full_text}],
            "total_pages": 1,
        }

    # PPTX processing removed - use PresentationLens service instead
    # See: https://github.com/michael-borck/presentation-lens

    async def _extract_from_text(self, content: bytes, filename: str) -> str:
        """Extract text from plain text or markdown file"""
        try:
            # Try different encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            return content.decode("utf-8", errors="replace")

        except Exception as e:
            raise ValueError(f"Failed to decode text file: {e!s}") from e

    async def _extract_from_json(self, content: bytes, filename: str) -> str:
        """Extract text from JSON file"""
        try:
            text = content.decode("utf-8")
            data = json.loads(text)

            # Extract text recursively from JSON structure
            extracted_texts: list[str] = []
            self._extract_text_from_json_recursive(data, extracted_texts)

            if not extracted_texts:
                raise ValueError("No text content found in JSON")

            return "\n\n".join(extracted_texts)

        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e!s}") from e
        except Exception as e:
            raise ValueError(f"Failed to process JSON: {e!s}") from e

    def _extract_text_from_json_recursive(self, obj: dict | list, texts: list[str]) -> None:
        """Recursively extract text from JSON object"""
        if isinstance(obj, dict):
            for _key, value in obj.items():
                if isinstance(value, str) and len(value.strip()) > 10:
                    texts.append(value.strip())
                elif isinstance(value, dict | list):
                    self._extract_text_from_json_recursive(value, texts)
        elif isinstance(obj, list):
            for item in obj:
                if isinstance(item, str) and len(item.strip()) > 10:
                    texts.append(item.strip())
                elif isinstance(item, dict | list):
                    self._extract_text_from_json_recursive(item, texts)

    def validate_file_size(self, content: bytes, max_size: int) -> bool:
        """Validate file size"""
        return len(content) <= max_size

    async def extract_metadata(
        self, content: bytes, content_type: str, filename: str
    ) -> dict[str, Any]:
        """
        Extract metadata from document

        Args:
            content: Raw file content
            content_type: MIME type of the file
            filename: Name of the file

        Returns:
            Dictionary containing extracted metadata
        """
        metadata = {"filename": filename, "size": len(content), "content_type": content_type}

        try:
            if content_type == "application/pdf":
                metadata.update(self._extract_pdf_metadata(content))
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                metadata.update(self._extract_docx_metadata(content))
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "application/vnd.ms-powerpoint",
            ]:
                metadata.update(self._extract_pptx_metadata(content))
            else:
                # For text files, just basic info
                if content_type.startswith("text/"):
                    try:
                        text = content.decode("utf-8")
                        metadata["lines"] = len(text.splitlines())
                        metadata["characters"] = len(text)
                    except UnicodeDecodeError:
                        metadata["encoding_error"] = True

        except Exception as e:
            metadata["extraction_error"] = str(e)

        return metadata

    def _extract_pdf_metadata(self, content: bytes) -> dict[str, Any]:
        """Extract metadata from PDF"""
        metadata: dict[str, Any] = {}

        if self.pdf_available:
            try:
                pdf_file = io.BytesIO(content)
                reader = PdfReader(pdf_file)

                metadata["pages"] = len(reader.pages)

                if reader.metadata:
                    doc_info = reader.metadata
                    metadata["title"] = str(doc_info.get("/Title", "")).strip()
                    metadata["author"] = str(doc_info.get("/Author", "")).strip()
                    metadata["subject"] = str(doc_info.get("/Subject", "")).strip()
                    metadata["creator"] = str(doc_info.get("/Creator", "")).strip()
                    metadata["producer"] = str(doc_info.get("/Producer", "")).strip()

                    # Handle dates
                    creation_date = doc_info.get("/CreationDate")
                    if creation_date:
                        metadata["creation_date"] = str(creation_date)

                    mod_date = doc_info.get("/ModDate")
                    if mod_date:
                        metadata["modification_date"] = str(mod_date)

            except Exception as e:
                metadata["pdf_metadata_error"] = str(e)

        return metadata

    def _extract_docx_metadata(self, content: bytes) -> dict[str, Any]:
        """Extract metadata from DOCX"""
        metadata: dict[str, Any] = {}

        if self.docx_available:
            try:
                docx_file = io.BytesIO(content)
                doc = Document(docx_file)

                # Count elements
                metadata["paragraphs"] = len(doc.paragraphs)
                metadata["tables"] = len(doc.tables)

                # Core properties
                if hasattr(doc, "core_properties"):
                    props = doc.core_properties
                    metadata["title"] = str(props.title or "")
                    metadata["author"] = str(props.author or "")
                    metadata["subject"] = str(props.subject or "")
                    metadata["keywords"] = str(props.keywords or "")
                    metadata["comments"] = str(props.comments or "")

                    if props.created:
                        metadata["creation_date"] = props.created.isoformat()
                    if props.modified:
                        metadata["modification_date"] = props.modified.isoformat()

            except Exception as e:
                metadata["docx_metadata_error"] = str(e)

        return metadata

    def _extract_pptx_metadata(self, content: bytes) -> dict[str, Any]:
        """Extract metadata from PPTX"""
        metadata: dict[str, Any] = {}

        if self.pptx_available:
            try:
                pptx_file = io.BytesIO(content)
                prs = Presentation(pptx_file)

                metadata["slides"] = len(prs.slides)

                # Core properties
                if hasattr(prs, "core_properties"):
                    props = prs.core_properties
                    metadata["title"] = str(props.title or "")
                    metadata["author"] = str(props.author or "")
                    metadata["subject"] = str(props.subject or "")
                    metadata["keywords"] = str(props.keywords or "")
                    metadata["comments"] = str(props.comments or "")

                    if props.created:
                        metadata["creation_date"] = props.created.isoformat()
                    if props.modified:
                        metadata["modification_date"] = props.modified.isoformat()

            except Exception as e:
                metadata["pptx_metadata_error"] = str(e)

        return metadata

    def infer_metadata_from_content(self, text: str, filename: str | None = None) -> dict[str, Any]:
        """
        Infer metadata from document content using pattern matching.

        Designed for corporate annual reports and sustainability documents.
        Extracts:
        - Year (fiscal year, report year)
        - Company name
        - Industry (based on keywords)
        - Document type (annual report, sustainability report, etc.)

        Args:
            text: Extracted document text
            filename: Optional filename for additional hints

        Returns:
            Dictionary matching InferredMetadata schema
        """
        result: dict[str, Any] = {
            "probable_year": None,
            "probable_company": None,
            "probable_industry": None,
            "document_type": None,
            "confidence_scores": {
                "year": 0.0,
                "company": 0.0,
                "industry": 0.0,
                "document_type": 0.0,
            },
            "extraction_notes": [],
        }

        # Use first ~10000 chars for faster processing (most metadata is at the start)
        sample_text = text[:10000]
        full_text_lower = text.lower()

        # === YEAR EXTRACTION ===
        result.update(self._infer_year(sample_text, filename))

        # === COMPANY NAME EXTRACTION ===
        result.update(self._infer_company(sample_text, filename))

        # === INDUSTRY DETECTION ===
        result.update(self._infer_industry(full_text_lower))

        # === DOCUMENT TYPE DETECTION ===
        result.update(self._infer_document_type(full_text_lower, filename))

        return result

    def _infer_year(self, sample_text: str, filename: str | None) -> dict[str, Any]:
        """Extract probable year from document"""
        result: dict[str, Any] = {
            "probable_year": None,
            "confidence_scores": {"year": 0.0},
            "extraction_notes": [],
        }

        year_candidates: dict[int, float] = {}

        # Pattern 1: "Annual Report 2023" or "2023 Annual Report"
        annual_report_pattern = r"(?:annual\s+report|sustainability\s+report|integrated\s+report)\s*(\d{4})|(\d{4})\s*(?:annual\s+report|sustainability\s+report|integrated\s+report)"
        for match in re.finditer(annual_report_pattern, sample_text, re.IGNORECASE):
            year_str = match.group(1) or match.group(2)
            year = int(year_str)
            if 1990 <= year <= 2030:
                year_candidates[year] = year_candidates.get(year, 0) + 0.9

        # Pattern 2: "Fiscal Year 2023" or "FY2023" or "FY 2023"
        fy_pattern = r"(?:fiscal\s+year|fy)\s*(\d{4})"
        for match in re.finditer(fy_pattern, sample_text, re.IGNORECASE):
            year = int(match.group(1))
            if 1990 <= year <= 2030:
                year_candidates[year] = year_candidates.get(year, 0) + 0.85

        # Pattern 3: "For the year ended December 31, 2023"
        year_ended_pattern = r"(?:for\s+the\s+)?year\s+ended?\s+\w+\s+\d{1,2},?\s*(\d{4})"
        for match in re.finditer(year_ended_pattern, sample_text, re.IGNORECASE):
            year = int(match.group(1))
            if 1990 <= year <= 2030:
                year_candidates[year] = year_candidates.get(year, 0) + 0.8

        # Pattern 4: Check filename for year
        if filename:
            filename_year_match = re.search(r"(20\d{2})", filename)
            if filename_year_match:
                year = int(filename_year_match.group(1))
                if 1990 <= year <= 2030:
                    year_candidates[year] = year_candidates.get(year, 0) + 0.7
                    result["extraction_notes"].append(f"Year {year} found in filename")

        # Pattern 5: Copyright years (lower confidence)
        copyright_pattern = r"(?:©|copyright)\s*(\d{4})"
        for match in re.finditer(copyright_pattern, sample_text, re.IGNORECASE):
            year = int(match.group(1))
            if 1990 <= year <= 2030:
                year_candidates[year] = year_candidates.get(year, 0) + 0.3

        # Select best year candidate
        if year_candidates:
            best_year = max(year_candidates, key=lambda y: year_candidates[y])
            result["probable_year"] = best_year
            result["confidence_scores"]["year"] = min(year_candidates[best_year], 1.0)

        return result

    def _infer_company(self, sample_text: str, filename: str | None) -> dict[str, Any]:
        """Extract probable company name from document"""
        result: dict[str, Any] = {
            "probable_company": None,
            "confidence_scores": {"company": 0.0},
            "extraction_notes": [],
        }

        company_candidates: dict[str, float] = {}

        # Pattern 1: "XYZ Corporation Annual Report" or "XYZ Inc. Sustainability Report"
        corp_patterns = [
            r"([A-Z][A-Za-z\s&]+(?:Corporation|Corp\.?|Inc\.?|Ltd\.?|Limited|PLC|plc|Group|Holdings|Company|Co\.?))\s+(?:Annual|Sustainability|Integrated)\s+Report",
            r"(?:Annual|Sustainability|Integrated)\s+Report\s+of\s+([A-Z][A-Za-z\s&]+(?:Corporation|Corp\.?|Inc\.?|Ltd\.?|Limited|PLC|plc|Group|Holdings|Company|Co\.?))",
        ]
        for pattern in corp_patterns:
            for match in re.finditer(pattern, sample_text):
                company = match.group(1).strip()
                if 3 <= len(company) <= 100:
                    company_candidates[company] = company_candidates.get(company, 0) + 0.85

        # Pattern 2: "About [Company Name]" section headers
        about_pattern = r"About\s+([A-Z][A-Za-z\s&]+(?:Corporation|Corp\.?|Inc\.?|Ltd\.?|Limited|PLC|plc|Group|Holdings)?)"
        for match in re.finditer(about_pattern, sample_text):
            company = match.group(1).strip()
            if 3 <= len(company) <= 100 and company.lower() not in ["this", "our", "the"]:
                company_candidates[company] = company_candidates.get(company, 0) + 0.6

        # Pattern 3: Extract from filename
        if filename:
            # Remove extension and year, clean up
            clean_name = re.sub(r"\.\w+$", "", filename)  # Remove extension
            clean_name = re.sub(r"[-_]?(20\d{2})[-_]?", " ", clean_name)  # Remove year
            clean_name = re.sub(
                r"[-_]?(annual|sustainability|report|integrated)[-_]?",
                " ",
                clean_name,
                flags=re.IGNORECASE,
            )
            clean_name = clean_name.strip()
            if len(clean_name) >= 3:
                company_candidates[clean_name] = company_candidates.get(clean_name, 0) + 0.5
                result["extraction_notes"].append(f"Company hint from filename: {clean_name}")

        # Select best company candidate
        if company_candidates:
            best_company = max(company_candidates, key=lambda c: company_candidates[c])
            result["probable_company"] = best_company
            result["confidence_scores"]["company"] = min(company_candidates[best_company], 1.0)

        return result

    def _infer_industry(self, full_text_lower: str) -> dict[str, Any]:
        """Detect probable industry from document content"""
        result: dict[str, Any] = {
            "probable_industry": None,
            "confidence_scores": {"industry": 0.0},
            "extraction_notes": [],
        }

        # Industry keyword mappings (keyword -> industry, weight)
        industry_keywords: dict[str, list[tuple[str, float]]] = {
            "Energy": [
                ("oil and gas", 0.9),
                ("petroleum", 0.8),
                ("natural gas", 0.7),
                ("drilling", 0.5),
                ("refinery", 0.7),
                ("upstream", 0.4),
                ("downstream", 0.4),
                ("barrel", 0.3),
                ("crude oil", 0.8),
                ("lng", 0.6),
                ("renewable energy", 0.7),
                ("solar power", 0.7),
                ("wind power", 0.7),
                ("hydroelectric", 0.7),
            ],
            "Financial Services": [
                ("banking", 0.8),
                ("investment management", 0.8),
                ("asset management", 0.7),
                ("wealth management", 0.7),
                ("insurance", 0.6),
                ("underwriting", 0.6),
                ("mortgage", 0.5),
                ("deposits", 0.4),
                ("loan portfolio", 0.7),
                ("credit risk", 0.6),
                ("capital adequacy", 0.7),
            ],
            "Technology": [
                ("software", 0.6),
                ("cloud computing", 0.8),
                ("artificial intelligence", 0.7),
                ("machine learning", 0.6),
                ("saas", 0.8),
                ("data center", 0.6),
                ("semiconductor", 0.8),
                ("cybersecurity", 0.7),
                ("digital transformation", 0.5),
            ],
            "Healthcare": [
                ("pharmaceutical", 0.9),
                ("clinical trial", 0.9),
                ("drug development", 0.8),
                ("patient", 0.4),
                ("hospital", 0.6),
                ("medical device", 0.8),
                ("fda approval", 0.9),
                ("therapeutic", 0.7),
                ("biotech", 0.8),
            ],
            "Consumer Goods": [
                ("retail", 0.6),
                ("consumer products", 0.8),
                ("brand portfolio", 0.7),
                ("fmcg", 0.9),
                ("packaged goods", 0.8),
                ("e-commerce", 0.5),
                ("store locations", 0.6),
            ],
            "Manufacturing": [
                ("manufacturing", 0.7),
                ("production facility", 0.7),
                ("supply chain", 0.5),
                ("factory", 0.6),
                ("industrial", 0.5),
                ("machinery", 0.6),
                ("assembly", 0.5),
                ("automotive", 0.7),
            ],
            "Utilities": [
                ("electric utility", 0.9),
                ("power generation", 0.8),
                ("transmission", 0.5),
                ("distribution network", 0.6),
                ("megawatt", 0.7),
                ("grid", 0.5),
                ("ratepayer", 0.8),
            ],
            "Mining & Metals": [
                ("mining", 0.8),
                ("ore", 0.7),
                ("extraction", 0.5),
                ("mineral", 0.6),
                ("gold", 0.5),
                ("copper", 0.5),
                ("iron ore", 0.8),
                ("smelting", 0.8),
            ],
            "Real Estate": [
                ("real estate", 0.9),
                ("property", 0.5),
                ("reit", 0.9),
                ("tenant", 0.6),
                ("occupancy rate", 0.8),
                ("square feet", 0.5),
                ("commercial property", 0.8),
            ],
            "Transportation": [
                ("airline", 0.9),
                ("aviation", 0.8),
                ("shipping", 0.7),
                ("logistics", 0.6),
                ("freight", 0.7),
                ("fleet", 0.5),
                ("passenger", 0.5),
                ("cargo", 0.6),
            ],
        }

        industry_scores: dict[str, float] = {}

        for industry, keywords in industry_keywords.items():
            score = 0.0
            for keyword, weight in keywords:
                count = full_text_lower.count(keyword)
                if count > 0:
                    # Diminishing returns for repeated mentions
                    score += weight * min(count, 10) / 10

            if score > 0:
                industry_scores[industry] = score

        if industry_scores:
            best_industry = max(industry_scores, key=lambda i: industry_scores[i])
            max_score = industry_scores[best_industry]

            # Normalize confidence (rough heuristic)
            confidence = min(max_score / 3.0, 1.0)

            if confidence >= 0.2:  # Only report if somewhat confident
                result["probable_industry"] = best_industry
                result["confidence_scores"]["industry"] = confidence

        return result

    def _infer_document_type(self, full_text_lower: str, filename: str | None) -> dict[str, Any]:
        """Detect document type from content"""
        result: dict[str, Any] = {
            "document_type": None,
            "confidence_scores": {"document_type": 0.0},
            "extraction_notes": [],
        }

        doc_type_patterns: dict[str, list[tuple[str, float]]] = {
            "Annual Report": [
                ("annual report", 0.9),
                ("form 10-k", 0.95),
                ("10-k", 0.8),
                ("fiscal year", 0.5),
                ("shareholders", 0.3),
                ("board of directors", 0.4),
            ],
            "Sustainability Report": [
                ("sustainability report", 0.95),
                ("esg report", 0.9),
                ("corporate responsibility", 0.7),
                ("environmental, social", 0.8),
                ("carbon footprint", 0.5),
                ("sustainable development goals", 0.7),
                ("sdg", 0.4),
                ("gri standards", 0.8),
                ("tcfd", 0.6),
            ],
            "Integrated Report": [
                ("integrated report", 0.95),
                ("integrated annual report", 0.95),
                ("value creation", 0.4),
                ("six capitals", 0.8),
                ("iirc", 0.9),
            ],
            "CSR Report": [
                ("csr report", 0.95),
                ("corporate social responsibility", 0.9),
                ("community investment", 0.5),
                ("social impact", 0.5),
            ],
            "Climate Report": [
                ("climate report", 0.95),
                ("climate-related", 0.6),
                ("net zero", 0.5),
                ("carbon neutral", 0.6),
                ("greenhouse gas", 0.5),
                ("scope 1", 0.6),
                ("scope 2", 0.6),
                ("scope 3", 0.6),
            ],
        }

        doc_type_scores: dict[str, float] = {}

        for doc_type, patterns in doc_type_patterns.items():
            score = 0.0
            for pattern, weight in patterns:
                if pattern in full_text_lower:
                    score += weight

            if score > 0:
                doc_type_scores[doc_type] = score

        # Also check filename
        if filename:
            filename_lower = filename.lower()
            for doc_type in doc_type_patterns:
                if doc_type.lower().replace(" ", "") in filename_lower.replace(" ", "").replace(
                    "-", ""
                ).replace("_", ""):
                    doc_type_scores[doc_type] = doc_type_scores.get(doc_type, 0) + 0.5
                    result["extraction_notes"].append("Document type hint from filename")

        if doc_type_scores:
            best_type = max(doc_type_scores, key=lambda t: doc_type_scores[t])
            max_score = doc_type_scores[best_type]

            # Normalize confidence
            confidence = min(max_score / 2.0, 1.0)

            if confidence >= 0.3:
                result["document_type"] = best_type
                result["confidence_scores"]["document_type"] = confidence

        return result
